"""file_extract — 附件真解析:PDF / Word(.docx)/ Excel(.xlsx)→ 纯文本。

files 桥(console 预览 /api/files/view)与 read_file 工具共用的**唯一**取文本入口 ——
解析产物走 CSV 同一条产线(截断/预览/交办语义由调用方沿用既有路径),不另起平行管线。

设计原则:
- **通用基建必借**:pypdf / python-docx / openpyxl(可选依赖 ``karvyloop[files]``),
  不自造解析器;没装 → ``ok=False`` + ``error="missing_dependency"`` + 明确安装提示,绝不崩。
- **解析器纪律(宁空勿毒)**:坏文件 / 伪造扩展名(magic 不符)/ 解析异常 →
  空文本 + 明确错误码,**绝不把二进制垃圾灌进上下文**。
- **截断有上限并明示**:提取文本封顶 ``max_chars``(默认与文本预览同一条 100KB 线),
  超限置 ``truncated=True``,由调用方向人/agent 明示。
"""
from __future__ import annotations

import io
from dataclasses import dataclass

#: 提取文本默认上限 —— 与 /api/files/view 的纯文本预览 100KB 上限同一条截断线。
MAX_EXTRACT_CHARS = 100_000

#: 缺依赖时的统一安装提示(用户可见文案由前端 i18n 出;这里是给 agent/日志的事实句)。
INSTALL_HINT = 'pip install "karvyloop[files]"'

#: 后缀 → 解析类别。CSV / 纯文本**不在此表**:它们本来就是文本,走原有 decode 路径。
_SUFFIX_KINDS = {".pdf": "pdf", ".docx": "docx", ".xlsx": "xlsx"}


@dataclass
class ExtractResult:
    """解析结果。``ok=False`` 时 ``text`` 恒为空串(宁空勿毒)。"""
    ok: bool
    kind: str = ""          # pdf | docx | xlsx
    text: str = ""
    truncated: bool = False
    error: str = ""         # "" | "missing_dependency" | "bad_file"
    hint: str = ""          # 人可读补充(缺哪个包 / 坏在哪),绝不含文件内容


def extract_kind(name: str) -> str | None:
    """按文件名后缀判定是否是需要解析的二进制文档格式;不是(含 CSV/文本)→ None。"""
    dot = name.rfind(".")
    return _SUFFIX_KINDS.get(name[dot:].lower()) if dot >= 0 else None


def extract_text(data: bytes, kind: str, *, max_chars: int = MAX_EXTRACT_CHARS) -> ExtractResult:
    """把 ``kind`` 格式的字节解析成纯文本(唯一对外入口)。

    伪造扩展名防线:先验 magic(PDF 头 / zip 容器 PK 头),不符直接 ``bad_file`` ——
    解析库还没碰到字节就被拒,坏文件伪装不进上下文。
    """
    if kind not in ("pdf", "docx", "xlsx"):
        return ExtractResult(False, kind, error="bad_file", hint=f"未知格式: {kind}")
    if not _magic_ok(data, kind):
        return ExtractResult(False, kind, error="bad_file",
                             hint="文件头与扩展名不符(伪造扩展名或已损坏)")
    if kind == "pdf":
        return _extract_pdf(data, max_chars)
    if kind == "docx":
        return _extract_docx(data, max_chars)
    return _extract_xlsx(data, max_chars)


# ---------------------------------------------------------------- internals

def _magic_ok(data: bytes, kind: str) -> bool:
    if kind == "pdf":
        return b"%PDF-" in data[:1024]   # 规范允许头前有少量杂字节
    return data[:2] == b"PK"             # docx / xlsx 都是 zip 容器


def _finish(kind: str, parts: list[str], max_chars: int, truncated: bool) -> ExtractResult:
    text = "\n".join(parts).replace("\x00", "").strip()
    if len(text) > max_chars:
        text, truncated = text[:max_chars], True
    return ExtractResult(True, kind, text, truncated=truncated)


def _extract_pdf(data: bytes, max_chars: int) -> ExtractResult:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ExtractResult(False, "pdf", error="missing_dependency",
                             hint=f"pypdf 未安装 — {INSTALL_HINT}")
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                return ExtractResult(False, "pdf", error="bad_file",
                                     hint="PDF 已加密,无法提取")
        parts: list[str] = []
        total, truncated = 0, False
        for i, page in enumerate(reader.pages):
            txt = (page.extract_text() or "").strip()
            if txt:
                parts.append(f"--- page {i + 1} ---\n{txt}")
                total += len(parts[-1])
            if total > max_chars:
                truncated = True
                break
        return _finish("pdf", parts, max_chars, truncated)
    except Exception as e:
        # 截断/损坏的 PDF、非法 xref 等 → 空 + 错误码,不吐半解析垃圾
        return ExtractResult(False, "pdf", error="bad_file", hint=type(e).__name__)


def _extract_docx(data: bytes, max_chars: int) -> ExtractResult:
    try:
        import docx
    except ImportError:
        return ExtractResult(False, "docx", error="missing_dependency",
                             hint=f"python-docx 未安装 — {INSTALL_HINT}")
    try:
        d = docx.Document(io.BytesIO(data))
        parts = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
        for tbl in d.tables:   # 表格按行铺开,制表符分列(与 CSV 分析同一消费形状)
            for row in tbl.rows:
                parts.append("\t".join(c.text.strip() for c in row.cells))
        return _finish("docx", parts, max_chars, truncated=False)
    except Exception as e:
        return ExtractResult(False, "docx", error="bad_file", hint=type(e).__name__)


def _extract_xlsx(data: bytes, max_chars: int) -> ExtractResult:
    try:
        import openpyxl
    except ImportError:
        return ExtractResult(False, "xlsx", error="missing_dependency",
                             hint=f"openpyxl 未安装 — {INSTALL_HINT}")
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        try:
            parts: list[str] = []
            total, truncated = 0, False
            for ws in wb.worksheets:
                parts.append(f"# sheet: {ws.title}")
                total += len(parts[-1])
                for row in ws.iter_rows(values_only=True):
                    parts.append(",".join("" if v is None else str(v) for v in row))
                    total += len(parts[-1]) + 1
                    if total > max_chars:
                        truncated = True
                        break
                if truncated:
                    break
            return _finish("xlsx", parts, max_chars, truncated)
        finally:
            wb.close()
    except Exception as e:
        return ExtractResult(False, "xlsx", error="bad_file", hint=type(e).__name__)
